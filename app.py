import streamlit as st
import tempfile
import os
import matplotlib.pyplot as plt
import pandas as pd
from utils.processor import analyze_audits
from docx import Document
from fpdf import FPDF

st.set_page_config(page_title="Audit Insights Pro", layout="wide")
st.title("📋 Audit Insights Pro")
st.markdown("Upload 2 to 4 audit reports in PDF format. The app will summarize each and compare findings.")

uploaded_files = st.file_uploader("Upload PDF audit reports", type="pdf", accept_multiple_files=True)

if uploaded_files and 2 <= len(uploaded_files) <= 4:
    with st.spinner("Analyzing audits with AI..."):
        temp_paths = []
        filenames = []
        for file in uploaded_files:
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
            temp_file.write(file.read())
            temp_file.close()
            temp_paths.append(temp_file.name)
            filenames.append(file.name)

        result = analyze_audits(temp_paths, filenames)

        summaries = result["summaries"]
        logos = result["logos"]
        themes = result["themes"]
        comparison = result["comparison"]
        learnings = result["learnings"]

        st.subheader("📝 Audit Summaries")
        for i, summary in enumerate(summaries):
            st.markdown(f"### 🧾 Audit: {filenames[i]}")
            if logos[i] and os.path.exists(logos[i]):
                st.image(logos[i], width=120)
            st.markdown(f"""
            <div style='background-color:#f9f9f9;padding:10px;border-radius:5px;border:1px solid #ddd;'>
            <pre>{summary}</pre>
            </div>
            """, unsafe_allow_html=True)

        # Visual chart
        st.subheader("📊 Common Audit Themes")
        theme_flat = [item for sublist in themes for item in sublist]
        theme_counts = pd.Series(theme_flat).value_counts()

        fig, ax = plt.subplots()
        theme_counts.plot(kind='barh', ax=ax)
        plt.xlabel("Frequency")
        plt.ylabel("Audit Themes")
        st.pyplot(fig)

        # Display comparison
        st.subheader("🔍 Comparison of Audits")
        st.text(comparison)

        # Display learnings
        st.subheader("💡 Key Learnings for Future Audits")
        st.text(learnings)

        # Download buttons
        def create_docx():
            doc = Document()
            doc.add_heading("Audit Insights Summary", 0)
            for i, summary in enumerate(summaries):
                doc.add_heading(f"Audit: {filenames[i]}", level=1)
                doc.add_paragraph(summary)
            doc.add_heading("Comparison", level=1)
            doc.add_paragraph(comparison)
            doc.add_heading("Learnings", level=1)
            doc.add_paragraph(learnings)
            path = os.path.join(tempfile.gettempdir(), "audit_summary.docx")
            doc.save(path)
            return path

        def create_pdf():
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", 'B', 16)
            pdf.cell(0, 10, "Audit Insights Summary", ln=True, align='C')
            pdf.set_font("Arial", '', 12)
            for i, summary in enumerate(summaries):
                pdf.ln(10)
                pdf.set_font("Arial", 'B', 12)
                pdf.cell(0, 10, f"Audit: {filenames[i]}", ln=True)
                if logos[i] and os.path.exists(logos[i]):
                    pdf.image(logos[i], w=50)
                pdf.set_font("Arial", '', 11)
                for line in summary.split("\n"):
                    pdf.multi_cell(0, 8, line)
            pdf.set_font("Arial", 'B', 12)
            pdf.ln(5)
            pdf.cell(0, 10, "Comparison:", ln=True)
            pdf.set_font("Arial", '', 11)
            pdf.multi_cell(0, 8, comparison)
            pdf.set_font("Arial", 'B', 12)
            pdf.cell(0, 10, "Learnings:", ln=True)
            pdf.set_font("Arial", '', 11)
            pdf.multi_cell(0, 8, learnings)
            pdf_path = os.path.join(tempfile.gettempdir(), "audit_summary.pdf")
            pdf.output(pdf_path)
            return pdf_path

        docx_file = create_docx()
        pdf_file = create_pdf()

        st.download_button("📥 Download DOCX Summary", open(docx_file, "rb"), file_name="audit_summary.docx")
        st.download_button("📥 Download PDF Summary", open(pdf_file, "rb"), file_name="audit_summary.pdf")
else:
    st.warning("Please upload 2 to 4 PDF audit reports.")
